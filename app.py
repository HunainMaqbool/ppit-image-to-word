import io
import subprocess
import sys
from pathlib import Path

import numpy as np
import pytesseract
import streamlit as st
from docx import Document
from PIL import Image

# Check if Tesseract is available
try:
    pytesseract.get_tesseract_version()
except Exception:
    st.error("⚠️ Tesseract OCR is not installed or not found in PATH.")
    st.info("On Streamlit Cloud, Tesseract should be installed via packages.txt")
    st.stop()


def run_ocr(image_bytes):
    pil_image = Image.open(io.BytesIO(image_bytes))

    gray = pil_image.convert("L")

    gray_np = np.array(gray)
    threshold = gray_np.mean()
    binary_np = (gray_np > threshold) * 255
    binary = Image.fromarray(binary_np.astype("uint8"))

    text = pytesseract.image_to_string(binary, lang="eng")
    return text


def text_to_docx(text, output_path=None):
    doc = Document()
    paragraphs = text.split("\n")
    for para_text in paragraphs:
        if para_text.strip():
            doc.add_paragraph(para_text.strip())
        else:
            doc.add_paragraph()
    if output_path:
        doc.save(output_path)
    else:
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer


st.set_page_config(page_title="Image to Word Converter", page_icon="📄", layout="centered")

st.title("📄 Image to Word Converter")
st.markdown("Upload an image (JPG/PNG) to extract text and download as Word document (.docx)")

uploaded_file = st.file_uploader("Choose an image file", type=["jpg", "jpeg", "png"])

if uploaded_file is not None:
    image_bytes = uploaded_file.read()
    st.image(image_bytes, caption="Uploaded Image", width="stretch")

    if st.button("Convert to Word Document", type="primary"):
        with st.spinner("Processing image with OCR..."):
            try:
                text = run_ocr(image_bytes)

                if text.strip():
                    st.success("OCR completed successfully!")
                    st.subheader("Extracted Text Preview:")
                    st.text_area("", text, height=200, disabled=True)

                    docx_buffer = text_to_docx(text)

                    st.download_button(
                        label="📥 Download Word Document (.docx)",
                        data=docx_buffer,
                        file_name=f"{Path(uploaded_file.name).stem}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                else:
                    st.warning("No text detected in the image. Please try a different image.")
            except Exception as e:
                st.error(f"Error processing image: {str(e)}")
                st.info("Make sure Tesseract OCR is available in the environment.")
else:
    st.info("👆 Please upload an image file to get started")
